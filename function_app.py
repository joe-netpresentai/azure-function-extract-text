import azure.functions as func
import logging
import base64
import io
import docx  # python-docx
from PyPDF2 import PdfReader  # PyPDF2

# This defines your function app and sets the default authorization level
app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)

# This decorator registers the function with an HTTP trigger
@app.route(route="ExtractTextFromDocument")
def ExtractTextFromDocument(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')

    try:
        # Get the JSON body from the request
        req_body = req.get_json()
        file_name = req_body.get('filename')
        base64_content = req_body.get('filecontent')

        if not file_name or not base64_content:
            return func.HttpResponse("Missing 'filename' or 'filecontent' in request body.", status_code=400)

        # Decode the Base64 string into bytes
        file_bytes = base64.b64decode(base64_content)
        file_stream = io.BytesIO(file_bytes)
        
        extracted_text = ""

        # Section: DOCX file handling
        if file_name.lower().endswith('.docx'):
            document = docx.Document(file_stream)
            for para in document.paragraphs:
                extracted_text += para.text + "\n"

        # Section: PDF file handling
        elif file_name.lower().endswith('.pdf'):
            reader = PdfReader(file_stream)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_text += page_text + "\n"

        else:
            return func.HttpResponse("Unsupported file type. Please use .docx or .pdf.", status_code=400)

        # Return the extracted text
        return func.HttpResponse(extracted_text, status_code=200, mimetype="text/plain; charset=utf-8")

    except Exception as e:
        logging.error(f"An error occurred: {e}")
        return func.HttpResponse(f"An error occurred during file processing: {str(e)}", status_code=500)